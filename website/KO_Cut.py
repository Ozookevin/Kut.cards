import tkinter as tk
import nltk
from textblob import TextBlob
from newspaper import Article
import requests
import json
from docx import Document
from docx.shared import Length
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.text.run import Font, Run
from nltk.tokenize import sent_tokenize, word_tokenize
import re
from datetime import datetime, date
import gensim 
import heapq
from urllib.parse import urlparse

# Things that need to be download

nltk.download('punkt')

def get_content(url):
    url = url 
    article = Article(url)
    article.download()
    article.parse()
    article.nlp()
    Title = article.title
    Authors = article.authors
    Publication_Date = article.publish_date
    clean_text = article.text
    cites = [Title, Authors, Publication_Date]
    
    return clean_text, cites, Title, Authors, Publication_Date
def sum_text(url):
    url = url
    response =['https://api.smmry.com/&SM_API_KEY=5BD00B32C7&SM_URL=']
    response.append(url)
    clean = ''.join(response)
    content = requests.get(clean)
    return content.text
    #print(' ')
    #print(' ')
    #print(content.text)

def jsonconvert(json_text):
    json_text = json_text
    jsonconvert = json.loads(json_text)
    return jsonconvert

def document(title,cites,content,clean_content):
    
    
    token_sum = sent_tokenize(content)
    token_full_content = sent_tokenize(clean_content)

    document = Document()
    
    body = document.add_paragraph(style = 'Body Text')
    for sentences in token_full_content:
        if any(sentences in s for s in token_sum):
            highlighted = body.add_run(sentences)
            highlighted.bold = True
            highlighted.underline = True
            highlighted.font.size = Pt(12)
        else:
            unhighlighted = body.add_run(sentences)
            unhighlighted.bold = False
            unhighlighted.underline = False
            unhighlighted.font.size = Pt(8)
    
   
    document.save("temp_content.docx")

    
#calls for program


url = input('Please Enter your URL-Content: ')


#clean_text, cites = get_content(url)
cites_content = get_content(url)
title = cites_content[2]
cites = str(cites_content[3])
Authors = str(cites_content[3])
pub = str(cites_content[4])
clean_content = cites_content[0]
json_text = sum_text(url)
json_list = jsonconvert(json_text)
content = json_list['sm_api_content']

#citeTag = cites_content[2]
citeDate = " "
pubDate = pub
citeQual = "No Qual Found please insert one!"
citePub = "  "
citeURL = url
citeFullname = Authors

#Clean Cite Data

def author(Authors):
    #tokens = nltk.word_tokenize(Authors)
   try:
       Authors = Authors.split()[1]
       Authors = re.sub(r'[^\w\s]','',Authors)
       return Authors
   except:
       return Authors
document(title,cites,content,clean_content)

def citeDate(pub):
    #pub = re.sub(r'[^\w\s]','',pub)
    try:
        pub = datetime.strptime(pub,'%Y-%m-%d %H:%M:%S')
        pub = pub.strftime("%Y")
        a_string = str(pub)
        a_length = len(a_string)
        c = a_string[a_length - 2: a_length]
        return c
    except:
        if 'None' in pub:
            pub1 = '21'
        else:
            pub1 = 'date_Error'
    return pub1

citeDate = citeDate(pub)
citeAuthor = author(Authors)
def cleanTag(content):
    #content_formated = re.sub('[^a-zA-Z]', ' ', content )
    #content = re.sub(r'\s+', ' ', content_formated)

    sentence_list = nltk.sent_tokenize(content)
    stopwords = nltk.corpus.stopwords.words('english')
    
    word_frequencies = {}
    for word in nltk.word_tokenize(content):
        if word not in stopwords:
            if word not in word_frequencies.keys():
                word_frequencies[word] = 1
        else:
    
            maximum_frequncy = max(word_frequencies.values())

    for word in word_frequencies.keys():
        word_frequencies[word] = (word_frequencies[word]/maximum_frequncy)

    sentence_scores = {}
    for sent in sentence_list:
        for word in nltk.word_tokenize(sent.lower()):
            if word in word_frequencies.keys():
                if len(sent.split(' ')) < 50:
                    if sent not in sentence_scores.keys():
                        sentence_scores[sent] = word_frequencies[word]
                    else:
                        sentence_scores[sent] += word_frequencies[word]
    
    summary_sentences = heapq.nlargest(1, sentence_scores, key=sentence_scores.get)
    summary = ' '.join(summary_sentences)
    sum1 = re.sub('[^a-zA-Z]', ' ', summary)
    return sum1
citeTag = cleanTag(clean_content)

def cleanPub(url):
    urlP = urlparse(url)
    urlp = urlparse(url).netloc
    urlp = ('.'.join(urlp.split('.')[1:]))
    urlp = (urlp[:-1])
    urlp = (urlp[:-1])
    urlp = (urlp[:-1])
    urlp = (urlp[:-1])
    return urlp

citePub = cleanPub(url)