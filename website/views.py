import ssl
ssl._create_default_https_context = ssl._create_unverified_context
from flask import Blueprint, render_template, send_from_directory, send_file
from flask.templating import render_template_string
from flask import request
from newspaper import Article
import pandas as pd
import requests
import json
from mailmerge import MailMerge
from datetime import date
from docx import Document
from docxcompose.composer import Composer
from docx import Document as Document_compose
import articleDateExtractor
import tldextract 
from docx import Document
from docx.shared import Length
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.text.run import Font, Run
import nltk
from website import Flask, create_app 
nltk.download()
from nltk.tokenize import sent_tokenize, word_tokenize
import re
import heapq
from datetime import datetime, date
import ctypes

views = Blueprint('views',__name__)
from bs4 import BeautifulSoup
import urllib.request,sys,time

app = create_app()

@views.route('/', methods =['POST','GET'])

def home():
    if request.method == "POST":
        url = request.form['nm']
        article = Article(url)
        article.download()
        article.parse()
        article.nlp()
        cite_title = article.title
        
        
        try:
            ext = tldextract.extract(url)
            publisher = ext.domain.title()
        except: 
            publisher = ''   
        try:
            page=requests.get(url)
            soup = BeautifulSoup(page.text, "html.parser")
            lastname = soup.find("meta",  {"property":"author"})
            lastname = lastname["content"] if lastname else None
            lastname1 = lastname.split()[1]
            lastname1 = re.sub(r'[^\w\s]','',lastname1)
            #lastname1 = str(lastname1)
            #lastname = lastname1
            # Using same method as above answer
            # author = author["content"] if author else None
        except:
            try:
                article = Article(url)
                article.download()
                article.parse()
                article.nlp()
                first_last = article.authors
                try:
                    lastname1 = first_last
                    lastname1 = lastname1.split()[1]
                    lastname1 = re.sub(r'[^\w\s]','',lastname1)
                    lastname1 = str(lastname1)
                except:
                    lastname1 = lastname
            except:
                lastname = publisher
                lastname1 = publisher

        try:
            cite_Publication_Date = article.publish_date
            date = datetime.strptime(cite_Publication_Date,'%Y-%m-%d %H:%M:%S')
            year = date.strftime("%Y")
            a_string = str(year)
            a_length = len(a_string)
            year = a_string[a_length - 2: a_length]
        except:
            cite_Publication_Date = "Last accessed 2021"
            year = '21'
            date = "Last assessed 2021"
           
        cite_clean_content = article.text
        
        template = "/Users/kevinozomaro/Documents/Website/Kut.cards/app/debate_temp.docx"


# calls to summery api
        response =['']
        response.append(url)
        clean_response_no_spaces = ''.join(response)
        content = requests.get(clean_response_no_spaces)
    # Convert Json into text
        json_response_text = content.text 
        jsonconvert = json.loads(json_response_text)
        clean_summery_content = jsonconvert['sm_api_content']

# create summery parsed document that will later be merged

        token_sum = sent_tokenize(clean_summery_content)
        token_full_content = sent_tokenize(cite_clean_content)

        document = Document()
    
        body = document.add_paragraph(style = 'Body Text')
        for sentences in token_full_content:
            if any(sentences in s for s in token_sum):
                highlighted = body.add_run(sentences)
                highlighted.bold = True
                highlighted.underline = True
                highlighted.font.size = Pt(12)
            else:
                unhighlighted = body.add_run(sentences)
                unhighlighted.bold = False
                unhighlighted.underline = False
                unhighlighted.font.size = Pt(8)
        
        document.save("temp_content.docx")

# Generate Card Tag:
        sentence_list = nltk.sent_tokenize(clean_summery_content)
        stopwords = nltk.corpus.stopwords.words('english')
        
        word_frequencies = {}
        for word in nltk.word_tokenize(clean_summery_content):
            if word not in stopwords:
                if word not in word_frequencies.keys():
                    word_frequencies[word] = 1
            else:
        
                maximum_frequncy = max(word_frequencies.values())

        for word in word_frequencies.keys():
            word_frequencies[word] = (word_frequencies[word]/maximum_frequncy)

        sentence_scores = {}
        for sent in sentence_list:
            for word in nltk.word_tokenize(sent.lower()):
                if word in word_frequencies.keys():
                    if len(sent.split(' ')) < 50:
                        if sent not in sentence_scores.keys():
                            sentence_scores[sent] = word_frequencies[word]
                        else:
                            sentence_scores[sent] += word_frequencies[word]
        
        summary_sentences = heapq.nlargest(1, sentence_scores, key=sentence_scores.get)
        summary = ' '.join(summary_sentences)
        sum1 = re.sub('[^a-zA-Z]', '', summary)
        cite_tag = summary

# applying MailMerge to input varables into word templete
        document = MailMerge(template)   
        document.merge(
            Tag = cite_tag,
            Author = lastname1,
            Date = year,
            Publication_date = date,
            Qual = cite_title,
            Author_full_name = lastname,
            content = cite_clean_content,
            Publication = publisher,
            URL = url
            )
        document.write('temp_cite.docx')
        



# merge documents together:
# combining documents together
        master = Document("temp_cite.docx")
        composer = Composer(master)
        doc1 = Document("temp_content.docx")
        composer.append(doc1)
        composer.save("app/website/Kut_io_card.docx")


        return send_file("Kut_io_card.docx", as_attachment= True)
    
    
    
    else:
        return render_template('Home.html')
